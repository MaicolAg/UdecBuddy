import os
import matplotlib
matplotlib.use('Agg')  # Usar el backend 'Agg' para evitar problemas con la GUI
import matplotlib.pyplot as plt

import io
import base64
#Django librerias
from django.views.generic.base import TemplateView
from django.shortcuts import render, redirect,render, get_object_or_404
from django.contrib.auth.mixins import LoginRequiredMixin
from django.contrib.auth.backends import ModelBackend
from django.contrib.auth import get_user_model,logout
from django.core.files.storage import FileSystemStorage
from django.conf import settings
from django.urls import reverse
from django.views import View
from django.http import HttpResponseRedirect, JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth.views import LoginView
from django.db.models import Count, Avg
from django.db.models.functions import TruncMonth
from django.contrib import messages
from django.core.exceptions import ValidationError
from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.utils.decorators import method_decorator

#Librerias internas
from .models import Usuarios, Archivos, ActividadChatbot, CalificacionRespuesta
from .forms import UsuariosForm, CambiarContrasenaForm,SinusuarioForm
from .chat import process_question
from .llm.utils import text_to_pinecone, cargar_archivos_nuevos, clean_files
#Langchain
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import Pinecone
from langchain.embeddings import HuggingFaceEmbeddings 
#Librerias de archivos
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
import pandas as pd
# Libreria .env 
from dotenv import load_dotenv



load_dotenv()

FILE_LIST = "archivos.txt"
PINECONE_ENV = "gcp-starter"
INDEX_NAME = 'test'




#Login y registro
class CustomLoginView(LoginView):
    def get_success_url(self):
        usuario = self.request.user  # Obtén el usuario que acaba de iniciar sesión
        if usuario.usuario == 'Administrador':
            return reverse('profileAdministrador')
        elif usuario.usuario == 'Directivo':
            return reverse('profileDirectivo')
        elif usuario.usuario == 'Estudiante':
            return reverse('profileEstudiante')

        else:
            return super().get_success_url()  # Si no es ninguno de los anteriores, usa la URL de éxito predeterminada

class EmailBackend(ModelBackend):
    def authenticate(self, request, username=None, password=None, **kwargs):
        UserModel = get_user_model()
        try:
            user = UserModel.objects.get(email=username)
        except UserModel.DoesNotExist:
            raise ValidationError("Usuario o contraseña incorrectos")

        if not user.check_password(password):
            raise ValidationError("Usuario o contraseña incorrectos")

        return user
        
class RegistroPageView(TemplateView):
    template_name =  "registro.html"
    def get(self, request):
        return render(request, self.template_name)
    def post(self, request):
        nombre = request.POST['nombre']
        apellido = request.POST['apellido']
        email = request.POST['email']
        password = request.POST['password']
        usuario = request.POST['usuario']
        semestre = request.POST['semestre']
        
        # Comprueba si el correo electrónico ya existe en la base de datos
        if Usuarios.objects.filter(email=email).exists():
            # Si el correo electrónico ya existe, renderiza la plantilla con un mensaje de error
            return render(request, self.template_name, {'error': 'Ese correo ya existe'})
        elif not email.endswith('@ucundinamarca.edu.co'):
            # Si el correo electrónico no termina con '@ucundinamarca.ed.co', renderiza la plantilla con un mensaje de error
            return render(request, self.template_name, {'error': 'El correo no es válido'})
        else:
            # Si el correo electrónico no existe y es válido, guarda el nuevo usuario
            us = Usuarios(nombre=nombre, apellido=apellido, email=email, usuario=usuario, semestre=semestre)
            us.set_password(password)
            us.save()
        return redirect('login')  # Redirige a la vista 'archivos
    
class LogoutPageView(TemplateView):
    def logout_view(request):
        logout(request)
        return redirect('login')

#Usuarios    
class EditarUsuarioView(View):
    def get(self, request, usuario_id):
        usuario = get_object_or_404(Usuarios, id=usuario_id)
        form = UsuariosForm(instance=usuario)
        return render(request, 'Administrativo/editar.html', {'form': form})

    def post(self, request, usuario_id):
        usuario = get_object_or_404(Usuarios, id=usuario_id)
        form = UsuariosForm(request.POST, instance=usuario)
        if form.is_valid():
            usuario = form.save(commit=False)  # No guardes el formulario todavía
            usuario.save()  # Ahora puedes guardar el formulario
            return redirect('table')
        else:
            return render(request, 'Administrativo/editar.html', {'form': form})

class EditarContrasenaUsuarioView(View):
    def get(self, request, usuario_id):
        usuario = get_object_or_404(Usuarios, id=usuario_id)
        form = CambiarContrasenaForm(instance=usuario)
        return render(request, 'Administrativo/editarcon.html', {'form': form})

    def post(self, request, usuario_id):
        usuario = get_object_or_404(Usuarios, id=usuario_id)
        form = CambiarContrasenaForm(request.POST, instance=usuario)
        if form.is_valid():
            usuario = form.save(commit=False)
            password = form.cleaned_data.get('password')
            usuario.set_password(password)
            usuario.save()# Ahora puedes guardar el formulario
            return redirect('table')
        else:
            return render(request, 'Administrativo/editarcon.html', {'form': form})


class EliminarUsuarioView(View):
    def get(self, request, usuario_id):
        usuario = get_object_or_404(Usuarios, id=usuario_id)
        usuario.delete()
        return redirect('table')

class RegistronnuevoView(TemplateView):
    template_name =  "Administrativo/registronuevo.html"
    def get(self, request):
        return render(request, self.template_name)
    def post(self, request):
        nombre = request.POST['nombre']
        apellido = request.POST['apellido']
        email = request.POST['email']
        password = request.POST['password']
        usuario = request.POST['usuario']
        semestre = request.POST['semestre']
        # Comprueba si el correo electrónico ya existe en la base de datos
        if Usuarios.objects.filter(email=email).exists():
            # Si el correo electrónico ya existe, renderiza la plantilla con un mensaje de error
            return render(request, self.template_name, {'error': 'Ese correo ya existe'})
        elif not email.endswith('@ucundinamarca.edu.co'):
            # Si el correo electrónico no termina con '@ucundinamarca.ed.co', renderiza la plantilla con un mensaje de error
            return render(request, self.template_name, {'error': 'El correo no es válido'})
        else:
            # Si el correo electrónico no existe y es válido, guarda el nuevo usuario
            us = Usuarios(nombre=nombre, apellido=apellido, email=email, usuario=usuario, semestre=semestre)
            us.set_password(password)
            us.save()
        return redirect('table')  # Redirige a la vista 'archivos'
    

#Archivos

def eliminar_nombre_archivo(path, archivo):
    if os.path.exists(path):
        with open(path, "r") as file:
            lines = file.readlines()
        with open(path, "w") as file:
            for line in lines:
                if line.strip("\n") != archivo:
                    file.write(line)

class EliminarArchivoView(View):
    def post(self, request, *args, **kwargs):
        archivo = request.POST.get('archivo')
        fs = FileSystemStorage(location='archivos')
        fs.delete(archivo)
        
        # Elimina el registro del archivo de la base de datos
        Archivos.objects.filter(nombre_archivo=archivo).delete()
        
        # Elimina el nombre del archivo del archivo de registro
        eliminar_nombre_archivo('archivos.txt', archivo)
        
        return HttpResponseRedirect(reverse('archivos'))
    
class GuardarTextoView(View):
    def post(self, request, *args, **kwargs):
        texto = request.POST.get('texto')
        lineas = [linea for linea in texto.split('\n') if linea.strip() != '']
        texto = '\n'.join(lineas)
        palabras = texto.split(',')[:6]
        nombre_archivo = '-'.join(palabras) 
        nombre_archivo = nombre_archivo.replace(' ', '_').replace('\r\n', '')
        nombre_archivo = nombre_archivo[:30]
        
        doc = SimpleDocTemplate('archivos/' + nombre_archivo + '.pdf', pagesize=letter)
        styles = getSampleStyleSheet()
        story = [Paragraph(texto, styles['BodyText'])]
        
        doc.build(story)
        
        Archivos.objects.create(nombre_archivo=nombre_archivo)
        
        return HttpResponseRedirect(reverse('archivos'))
    

class SubirArchivosView(LoginRequiredMixin, View):
    def post(self, request, *args, **kwargs):
        # Obtiene la lista de archivos en la carpeta 'archivos'
        archivos = os.listdir('archivos')
        extensiones_permitidas = ['.pdf', '.docx', '.txt']
        
        for nombre_archivo in archivos:
            # Verifica la extensión del archivo
            extension = os.path.splitext(nombre_archivo)[1].lower()
            if extension not in extensiones_permitidas:
                return HttpResponse("Solo se permiten archivos con las extensiones .pdf, .docx o .txt.", status=400)
        
        # Llama a la función para cargar solo los archivos nuevos
        cargar_archivos_nuevos()
        
        return HttpResponseRedirect(reverse('archivos'))
    
class EliminarArchivosView(View,LoginRequiredMixin):
    def post(self, request, *args, **kwargs):
        # Obtiene la lista de archivos en la carpeta 'archivos '
        archivos = os.listdir('archivos')
        clean_files()

        return HttpResponseRedirect(reverse('archivos'))


#Archivos Directivo

class EliminarArchivoDirectivoView(View):
    def post(self, request, *args, **kwargs):
        archivo = request.POST.get('archivo')
        fs = FileSystemStorage(location='archivos')
        fs.delete(archivo)
        
        # Elimina el registro del archivo de la base de datos
        Archivos.objects.filter(nombre_archivo=archivo).delete()
        
        return HttpResponseRedirect(reverse('archivosDirectivo'))
    
class GuardarTextoDirectivoView(View):
    def post(self, request, *args, **kwargs):
        texto = request.POST.get('texto')
        lineas = [linea for linea in texto.split('\n') if linea.strip() != '']
        texto = '\n'.join(lineas)
        palabras = texto.split(',')[:6]
        nombre_archivo = '-'.join(palabras) 
        nombre_archivo = nombre_archivo.replace(' ', '_').replace('\r\n', '')
        nombre_archivo = nombre_archivo[:30]
        
        doc = SimpleDocTemplate('archivos/' + nombre_archivo + '.pdf', pagesize=letter)
        styles = getSampleStyleSheet()
        story = [Paragraph(texto, styles['BodyText'])]
        
        doc.build(story)
        
        Archivos.objects.create(nombre_archivo=nombre_archivo)
        
        return HttpResponseRedirect(reverse('archivosDirectivo'))
    
class SubirArchivosDirectivoView(View,LoginRequiredMixin):
    def post(self, request, *args, **kwargs):
        # Obtiene la lista de archivos en la carpeta 'archivos'
        archivos = os.listdir('archivos')

        # Crea los embeddings y los sube a Pinecone
        for nombre_archivo in archivos:
            with open(os.path.join('archivos', nombre_archivo), 'rb') as f:
                pdf_content = f.read()

            if pdf_content:  # Comprueba si el archivo está vacío
                text_to_pinecone(pdf_content, nombre_archivo)

        return HttpResponseRedirect(reverse('archivosDirectivo'))
    
class EliminarArchivosDIrectivoView(View,LoginRequiredMixin):
    def post(self, request, *args, **kwargs):
        # Obtiene la lista de archivos en la carpeta 'archivos '
        archivos = os.listdir('archivos')
        clean_files()

        return HttpResponseRedirect(reverse('archivosDirectivo'))
   
#Administrativo

class TableAdministrativoPageView(LoginRequiredMixin,TemplateView):
    template_name = "Administrativo/table.html"
    def get(self, request):
        usuarios = Usuarios.objects.all()
        return render(request, self.template_name, {'usuarios': usuarios})
    
class ArchivosPageView(LoginRequiredMixin, TemplateView):
    template_name = "Administrativo/archivos.html"
    def get(self, request):
        # Obtén una lista de los nombres de los archivos en la carpeta 'archivos'
        archivos = os.listdir(settings.MEDIA_ROOT)
        # Ordena los archivos por fecha de modificación en orden descendente
        archivos.sort(key=lambda x: os.path.getmtime(os.path.join(settings.MEDIA_ROOT, x)), reverse=True)
        return render(request, self.template_name, {'archivos': archivos})
    
    def post(self, request):
        archivo_cargado = request.FILES['file-upload'] if 'file-upload' in request.FILES else None
        if archivo_cargado is not None:
            
            extension = archivo_cargado.name.rsplit('.', 1)[1]
            if extension.lower() == 'pdf':
                # Si el archivo ya es un PDF, guárdalo directamente
                fs = FileSystemStorage(location='archivos')  # especifica la ubicación donde se guardarán los archivos
                nombre = fs.save(archivo_cargado.name, archivo_cargado)
            else:
                # Crea un nuevo nombre de archivo con la extensión .pdf
                nombre = archivo_cargado.name.rsplit('.', 1)[0] + '.pdf'
                
                # Crea un objeto de canvas de reportlab
                c = canvas.Canvas('archivos/' + nombre)

                # Escribe el contenido del archivo en el PDF
                textobject = c.beginText(40, 800)
                
                # Comprueba la extensión del archivo para determinar cómo leerlo
                extension = archivo_cargado.name.rsplit('.', 1)[1]
                if extension in ['doc', 'docx']:
                    doc = Document(archivo_cargado)
                    for para in doc.paragraphs:
                        textobject.textLine(para.text)
                elif extension in ['xls', 'xlsx']:
                    df = pd.read_excel(archivo_cargado)
                    for row in df.iterrows():
                        textobject.textLine(', '.join(str(x) for x in row))
                else:  # asume que es un archivo de texto si no es Word o Excel
                    contenido = archivo_cargado.read().decode('utf-8')
                    for line in contenido.splitlines():
                        textobject.textLine(line)
            
                c.drawText(textobject)

                # Guarda el PDF
                c.save()

            # Guarda la información del archivo en la base de datos
            Archivos.objects.create(nombre_archivo=nombre)

            return HttpResponseRedirect(reverse('archivos'))  # redirige a la misma vista
        else:
            return render(request, self.template_name)
        
class DatosAdministrativoPageView(LoginRequiredMixin,TemplateView):
    template_name = "Administrativo/datos.html"
    def get(self, request):
        archivos = Archivos.objects.all()
        return render(request, self.template_name, {'archivos': archivos})

class ChatAdministradorPageView(LoginRequiredMixin, TemplateView):
    template_name = "Administrativo/chat.html"


class ChatMessageView(TemplateView, LoginRequiredMixin):
    @staticmethod
    def get(request):
        user_message = request.GET.get('message')
        response = process_question(user_message)

        # Registrar la actividad del chatbot
        actividad = ActividadChatbot(usuario=request.user)
        actividad.save()

        return JsonResponse({'response': response, 'actividad_id': actividad.id})

    
class PerfilAdministradorView(TemplateView):
    template_name = "Administrativo/profile.html"

    @method_decorator(login_required)  # Verifica la autenticación
    def get(self, request):
        usuario = request.user
        form = UsuariosForm(instance=usuario)
        return render(request, 'Administrativo/profile.html', {'form': form})

    @method_decorator(login_required)  # Verifica la autenticación
    def post(self, request):
        usuario = request.user
        form = UsuariosForm(request.POST, instance=usuario)
        if form.is_valid():
            usuario = form.save(commit=False)
            password = form.cleaned_data.get('password')
            if password:  # Si se proporciona una nueva contraseña
                usuario.set_password(password)
            else:
                # Mantén la contraseña existente
                usuario.password = usuario.password  # No es necesario guardar aquí
            
            usuario.save()
            return render(request, 'Administrativo/profile.html', {'form': form})  # Redirige al perfil
        else:
            return render(request, 'Administrativo/profile.html', {'form': form})
        
        
class CambiarContrasenaAdminView(TemplateView):
    template_name = "cambiar_contrasena.html"

    @method_decorator(login_required)  # Verifica la autenticación
    def get(self, request):
        usuario = request.user
        form = CambiarContrasenaForm(instance=usuario)
        return render(request, 'Administrativo/cambiarcon_admin.html', {'form': form})

    @method_decorator(login_required)  # Verifica la autenticación
    def post(self, request):
        usuario = request.user
        form = CambiarContrasenaForm(request.POST, instance=usuario)
        if form.is_valid():
            usuario = form.save(commit=False)
            password = form.cleaned_data.get('password')
            usuario.set_password(password)
            usuario.save()
            logout(request)
            return redirect('login')# Cierra la sesión del usuario
        else:
            return render(request, 'Administrativo/profile.html', {'form': form})
   
   
        
#Directivo


class TableDirectivoPageView(LoginRequiredMixin,TemplateView):
    template_name = "Directivo/table.html"
    def get(self, request):
        usuarios = Usuarios.objects.all()
        return render(request, self.template_name, {'usuarios': usuarios})

class ArchivosDirectivoPageView(LoginRequiredMixin, TemplateView):
    template_name = "Directivo/archivos.html"
    def get(self, request):
        # Obtén una lista de los nombres de los archivos en la carpeta 'archivos'
        archivos = os.listdir(settings.MEDIA_ROOT)
        # Ordena los archivos por fecha de modificación en orden descendente
        archivos.sort(key=lambda x: os.path.getmtime(os.path.join(settings.MEDIA_ROOT, x)), reverse=True)
        return render(request, self.template_name, {'archivos': archivos})
    
    def post(self, request):
        archivo_cargado = request.FILES['file-upload'] if 'file-upload' in request.FILES else None
        if archivo_cargado is not None:
            
            extension = archivo_cargado.name.rsplit('.', 1)[1]
            if extension.lower() == 'pdf':
                # Si el archivo ya es un PDF, guárdalo directamente
                fs = FileSystemStorage(location='archivos')  # especifica la ubicación donde se guardarán los archivos
                nombre = fs.save(archivo_cargado.name, archivo_cargado)
            else:
                # Crea un nuevo nombre de archivo con la extensión .pdf
                nombre = archivo_cargado.name.rsplit('.', 1)[0] + '.pdf'
                
                # Crea un objeto de canvas de reportlab
                c = canvas.Canvas('archivos/' + nombre)

                # Escribe el contenido del archivo en el PDF
                textobject = c.beginText(40, 800)
                
                # Comprueba la extensión del archivo para determinar cómo leerlo
                extension = archivo_cargado.name.rsplit('.', 1)[1]
                if extension in ['doc', 'docx']:
                    doc = Document(archivo_cargado)
                    for para in doc.paragraphs:
                        textobject.textLine(para.text)
                elif extension in ['xls', 'xlsx']:
                    df = pd.read_excel(archivo_cargado)
                    for row in df.iterrows():
                        textobject.textLine(', '.join(str(x) for x in row))
                else:  # asume que es un archivo de texto si no es Word o Excel
                    contenido = archivo_cargado.read().decode('utf-8')
                    for line in contenido.splitlines():
                        textobject.textLine(line)
            
                c.drawText(textobject)

                # Guarda el PDF
                c.save()

            # Guarda la información del archivo en la base de datos
            Archivos.objects.create(nombre_archivo=nombre)

            return HttpResponseRedirect(reverse('archivosDirectivo'))  # redirige a la misma vista
        else:
            return render(request, self.template_name)
        

class ChatDirectivoPageView(LoginRequiredMixin, TemplateView):
    template_name = "Directivo/chat.html"


class ChatDirectivoMessageView(TemplateView, LoginRequiredMixin):
    @staticmethod
    def get(request):
        user_message = request.GET.get('message')
        response = process_question(user_message)
        return JsonResponse({'response': response})


class PerfilDirectivoView(TemplateView, LoginRequiredMixin):
    template_name = "Directivo/profile.html"
    def get(self, request):
        usuario = request.user  # Obtén el usuario actualmente autenticado
        form = SinusuarioForm(instance=usuario)  # Asegúrate de tener un formulario para Usuarios
        return render(request, 'Directivo/profile.html', {'form': form})

    def post(self, request):
        usuario = request.user  # Obtén el usuario actualmente autenticado
        form = SinusuarioForm(request.POST, instance=usuario)
        if form.is_valid():
            form.save()
            return redirect('profileDirectivo')
        else:
            return render(request, 'Directivo/profile.html', {'form': form})

class CambiarContrasenaDireView(TemplateView):
    template_name = "Directivo/cambiarcon_dire.html"

    @method_decorator(login_required)  # Verifica la autenticación
    def get(self, request):
        usuario = request.user
        form = CambiarContrasenaForm(instance=usuario)# Asegúrate de tener un formulario para Usuarios
        return render(request, 'Directivo/cambiarcon_dire.html', {'form': form})

    @method_decorator(login_required)  # Verifica la autenticación
    def post(self, request):
        usuario = request.user
        form = CambiarContrasenaForm(request.POST, instance=usuario)
        if form.is_valid():
            usuario = form.save(commit=False)
            password = form.cleaned_data.get('password')
            usuario.set_password(password)
            usuario.save()
            logout(request)
            return redirect('login')# Cierra la sesión del usuario
        else:
            return render(request, 'Directivo/cambiarcon_dire.html', {'form': form})        
#Estudiante

class ChatEstudiantePageView(LoginRequiredMixin, TemplateView):
    template_name = "Estudiante/chat.html"
    
class ChatEstudianteMessageView(TemplateView, LoginRequiredMixin):
    @staticmethod
    def get(request):
        user_message = request.GET.get('message')
        response = process_question(user_message)
        return JsonResponse({'response': response})

class PerfilEstudiantePageView(TemplateView, LoginRequiredMixin):
    template_name = "Estudiante/profile.html"
    def get(self, request):
        usuario = request.user  # Obtén el usuario actualmente autenticado
        form = SinusuarioForm(instance=usuario)  # Asegúrate de tener un formulario para Usuarios
        return render(request, 'Estudiante/profile.html', {'form': form})


    def post(self, request):
        usuario = request.user  # Obtén el usuario actualmente autenticado
        form = SinusuarioForm(request.POST, instance=usuario)
        if form.is_valid():
            form.save()

            return redirect('profileEstudiante')
        else:
            return render(request, 'Estudiante/profile.html', {'form': form})

class CambiarContrasenaEstuView(TemplateView):
    template_name = "Estudiante//cambiarcon_estu.html"

    @method_decorator(login_required)  # Verifica la autenticación
    def get(self, request):
        usuario = request.user
        form = CambiarContrasenaForm(instance=usuario)# Asegúrate de tener un formulario para Usuarios
        return render(request, 'Estudiante//cambiarcon_estu.html', {'form': form})

    @method_decorator(login_required)  # Verifica la autenticación
    def post(self, request):
        usuario = request.user
        form = CambiarContrasenaForm(request.POST, instance=usuario)
        if form.is_valid():
            usuario = form.save(commit=False)
            password = form.cleaned_data.get('password')
            usuario.set_password(password)
            usuario.save()
            logout(request)
            return redirect('login')# Cierra la sesión del usuario
        else:
            return render(request, 'Estudiante//cambiarcon_estu.html', {'form': form})  
        

# Calificaciones bot

@csrf_exempt
def calificar_respuesta(request):
    if request.method == 'POST':
        calificacion = request.POST.get('calificacion')

        # Registrar la calificación de la respuesta
        calificacion_respuesta = CalificacionRespuesta(calificacion=calificacion)
        calificacion_respuesta.save()

        return JsonResponse({'success': True})




class EstadisticasPageView(TemplateView):
    template_name = 'Administrativo/estadisticas.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        
        # Obtener el número de interacciones por mes
        interacciones_por_mes = (ActividadChatbot.objects
                                 .annotate(mes=TruncMonth('fecha_registro'))
                                 .values('mes')
                                 .annotate(conteo=Count('id'))
                                 .order_by('mes'))
        
        # Preparar los datos para la gráfica de interacciones
        meses_interacciones = [interaccion['mes'].strftime('%m') for interaccion in interacciones_por_mes]
        conteos_interacciones = [interaccion['conteo'] for interaccion in interacciones_por_mes]
        
        # Crear la gráfica de barras para interacciones
        plt.figure(figsize=(5,5))  # Ajustar el tamaño de la gráfica
        plt.bar(meses_interacciones, conteos_interacciones, color='skyblue')
        plt.title('Número de Interacciones por Mes')
        plt.xlabel('Mes')
        plt.ylabel('Número de Interacciones')
        plt.xticks(rotation=45)
        plt.grid(True)
        
        # Guardar la gráfica en un buffer
        buffer = io.BytesIO()
        plt.savefig(buffer, format='png')
        buffer.seek(0)
        image_png = buffer.getvalue()
        buffer.close()
        
        # Codificar la imagen en base64
        graphic_interacciones = base64.b64encode(image_png)
        graphic_interacciones = graphic_interacciones.decode('utf-8')
        
        context['graphic_interacciones'] = graphic_interacciones
        
        # Obtener la calificación promedio por mes
        calificaciones_por_mes = (CalificacionRespuesta.objects
                                  .annotate(mes=TruncMonth('fecha_calificacion'))
                                  .values('mes')
                                  .annotate(promedio=Avg('calificacion'))
                                  .order_by('mes'))
        
        # Preparar los datos para la gráfica de calificaciones
        meses_calificaciones = [calificacion['mes'].strftime('%m') for calificacion in calificaciones_por_mes]
        promedios_calificaciones = [calificacion['promedio'] for calificacion in calificaciones_por_mes]
        
        # Crear la gráfica de barras para calificaciones
        plt.figure(figsize=(5, 5))  # Ajustar el tamaño de la gráfica
        plt.bar(meses_calificaciones, promedios_calificaciones, color='lightcoral')
        plt.title('Calificación Promedio por Mes')
        plt.xlabel('Mes')
        plt.ylabel('Calificación Promedio')
        plt.xticks(rotation=45)
        plt.ylim(1, 5)  # Asegurarse de que el eje Y vaya de 1 a 5
        plt.grid(True)
        
        # Guardar la gráfica en un buffer
        buffer = io.BytesIO()
        plt.savefig(buffer, format='png')
        buffer.seek(0)
        image_png = buffer.getvalue()
        buffer.close()
        
        # Codificar la imagen en base64
        graphic_calificaciones = base64.b64encode(image_png)
        graphic_calificaciones = graphic_calificaciones.decode('utf-8')
        
        context['graphic_calificaciones'] = graphic_calificaciones
        
        return context
